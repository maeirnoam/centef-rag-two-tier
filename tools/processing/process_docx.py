"""
DOCX processing tool for CENTEF RAG system.
Chunks DOCX files by paragraphs/sections and writes to GCS.
"""
import argparse
import logging
import os
import sys
from pathlib import Path
from typing import List, Tuple

from docx import Document
from google.cloud import storage

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from shared.schemas import Chunk, ChunkMetadata, ChunkAnchor, write_chunks_to_jsonl
from shared.manifest import get_manifest_entry, update_manifest_entry, DocumentStatus

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Environment variables
PROJECT_ID = os.getenv("PROJECT_ID")
TARGET_BUCKET = os.getenv("TARGET_BUCKET", "centef-rag-chunks")


def download_from_gcs(gcs_path: str, local_path: str) -> str:
    """
    Download file from GCS if needed.
    
    Args:
        gcs_path: GCS path (gs://bucket/path) or local path
        local_path: Where to save locally
    
    Returns:
        Local file path
    """
    if not gcs_path.startswith("gs://"):
        return gcs_path
    
    logger.info(f"Downloading {gcs_path} to {local_path}")
    
    # Parse GCS path
    parts = gcs_path.replace("gs://", "").split("/", 1)
    bucket_name = parts[0]
    blob_path = parts[1]
    
    client = storage.Client(project=PROJECT_ID)
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(blob_path)
    blob.download_to_filename(local_path)
    
    return local_path


def extract_docx_sections(docx_path: str, chunk_size: int = 5) -> List[Tuple[str, str]]:
    """
    Extract text from DOCX file, grouped into sections.
    
    Sections are created by grouping paragraphs. Headings start new sections.
    
    Args:
        docx_path: Path to DOCX file (local or GCS)
        chunk_size: Number of paragraphs per chunk (if no headings)
    
    Returns:
        List of tuples (section_name, text_content)
    """
    logger.info(f"Extracting text from DOCX: {docx_path}")
    
    # Download from GCS if needed
    if docx_path.startswith("gs://"):
        local_path = f"/tmp/{os.path.basename(docx_path)}"
        docx_path = download_from_gcs(docx_path, local_path)
    
    sections = []
    
    try:
        doc = Document(docx_path)
        
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
        
        current_section = "Introduction"
        current_content = []
        paragraph_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a heading (based on style)
            is_heading = para.style.name.startswith('Heading')
            
            if is_heading:
                # Save previous section if it has content
                if current_content:
                    sections.append((current_section, "\n\n".join(current_content)))
                
                # Start new section
                current_section = text
                current_content = []
                paragraph_count = 0
            else:
                # Add to current section
                current_content.append(text)
                paragraph_count += 1
                
                # If we've accumulated enough paragraphs and no heading was found, create a chunk
                if paragraph_count >= chunk_size and current_section == "Introduction":
                    section_name = f"Section {len(sections) + 1}"
                    sections.append((section_name, "\n\n".join(current_content)))
                    current_content = []
                    paragraph_count = 0
        
        # Add final section if it has content
        if current_content:
            sections.append((current_section, "\n\n".join(current_content)))
        
        logger.info(f"Extracted {len(sections)} sections from DOCX")
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}", exc_info=True)
        raise
    
    return sections


def process_docx(source_id: str, input_path: str) -> str:
    """
    Process a DOCX file into section-based chunks.
    
    Args:
        source_id: The source_id from manifest
        input_path: Path to DOCX file (local or GCS)
    
    Returns:
        Path to output JSONL file in GCS
    """
    logger.info(f"Processing DOCX for source_id={source_id}, input={input_path}")
    
    # Get manifest entry
    entry = get_manifest_entry(source_id)
    if not entry:
        raise ValueError(f"Manifest entry not found for source_id={source_id}")
    
    # Extract sections
    sections = extract_docx_sections(input_path)
    
    # Create chunks
    chunks = []
    for i, (section_name, text) in enumerate(sections):
        chunk_id = f"{source_id}_section_{i}"
        
        metadata = ChunkMetadata(
            id=chunk_id,
            source_id=source_id,
            filename=entry.filename,
            title=entry.title,
            mimetype=entry.mimetype,
            author=entry.author,
            organization=entry.organization,
            date=entry.date,
            publisher=entry.publisher,
            tags=entry.tags
        )
        
        anchor = ChunkAnchor(section=section_name)
        
        chunk = Chunk(
            metadata=metadata,
            anchor=anchor,
            content=text,
            chunk_index=i
        )
        
        chunks.append(chunk)
    
    logger.info(f"Created {len(chunks)} section-based chunks")
    
    # Write to GCS
    output_path = f"gs://{TARGET_BUCKET}/data/{source_id}.jsonl"
    
    # Write locally then upload
    local_output = f"/tmp/{source_id}.jsonl"
    write_chunks_to_jsonl(chunks, local_output)
    
    # Upload to GCS
    client = storage.Client(project=PROJECT_ID)
    bucket = client.bucket(TARGET_BUCKET.replace("gs://", ""))
    blob = bucket.blob(f"data/{source_id}.jsonl")
    blob.upload_from_filename(local_output)
    
    logger.info(f"Uploaded chunks to {output_path}")
    
    # Update manifest status
    update_manifest_entry(source_id, {
        "status": DocumentStatus.PENDING_SUMMARY,
        "data_path": output_path
    })
    
    logger.info(f"Updated manifest status to {DocumentStatus.PENDING_SUMMARY}")
    
    return output_path


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(description="Process DOCX into chunks")
    parser.add_argument("--source-id", required=True, help="Source ID from manifest")
    parser.add_argument("--input", required=True, help="Path to input DOCX file")
    
    args = parser.parse_args()
    
    try:
        output_path = process_docx(args.source_id, args.input)
        print(f"Successfully processed DOCX. Output: {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error processing DOCX: {e}", exc_info=True)
        return 1


if __name__ == "__main__":
    sys.exit(main())
